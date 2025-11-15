"""
ChatGPT integration for generating weekly content:
- 7 topics for voice journaling
- Shadowing scripts (two scripts per week)
- Weekly speaking prompt with 5 words

Background Context:
This tool is part of a 24-month speaking practice roadmap designed to build English speaking fluency
and confidence in professional contexts. The overall goal is to develop clear, articulate communication
skills for product management and business contexts.

Phase 1: Daily Speaking Habits (0-6 months)
- Objective: Build consistency, real-time speaking flow, and natural delivery
- Focus: Establishing daily speaking habits and building foundational fluency

Activities:
1. Voice Journaling: Record daily voice notes (2-3 mins), using 7 generated topics per week
2. Shadowing Practice: Daily practice with audio scripts to improve pronunciation and natural speaking rhythm
3. Weekly Speaking Prompt: Practice speaking about PM concepts (3-5 mins), using 5 generated words

Content should:
- Help build professional thinking and communication skills
- Focus on practical, useful vocabulary
- Progress appropriately over weeks/months
- Be engaging and motivating for daily practice
"""

import os
import csv
import random
from config import OPENAI_API_KEY

# Try to import python-docx for reading resume
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Load prompts from prompts.py if it exists, otherwise use defaults
try:
    from prompts import (
        BACKGROUND_CONTEXT,
        VOICE_JOURNALING_TOPICS_SYSTEM_PROMPT,
        VOICE_JOURNALING_TOPICS_USER_PROMPT,
        WEEKLY_PROMPT_WORDS_SYSTEM_PROMPT_EXTRA,
        WEEKLY_PROMPT_WORDS_USER_PROMPT,
        SHADOWING_SCRIPT_SYSTEM_PROMPT_EXTRA,
        SHADOWING_SCRIPT_USER_PROMPT,
        WEEKLY_PROMPT_SYSTEM_PROMPT_EXTRA,
        WEEKLY_PROMPT_USER_PROMPT,
        WEEKLY_PROMPT_BEST_ANSWER_SYSTEM_PROMPT,
        WEEKLY_PROMPT_BEST_ANSWER_USER_PROMPT,
        WEEKLY_PROMPT_BEST_ANSWER_HINTS_PROMPT,
    )
except ImportError:
    # Default prompts (used if prompts.py doesn't exist)
    BACKGROUND_CONTEXT = """You are helping a learner who is following a 24-month speaking practice roadmap to build professional English communication skills. 

Context:
- Goal: Develop fluency and confidence in professional speaking contexts (currently in Phase 1: 0-6 months)
- Phase 1 Focus: Daily Speaking Habits - building consistency, real-time speaking flow, and natural delivery
- Learning approach: Progressive weekly practice with three main activities
- Target: Professional contexts that require strong English communication skills

The learner is committed to building speaking fluency and the ability to articulate professional concepts clearly. 
Content should be practical, engaging, and progressively challenging. Each week builds on previous practice."""
    
    # Default prompt components (can be customized in prompts.py)
    VOICE_JOURNALING_TOPICS_SYSTEM_PROMPT = """You are helping a learner practice daily voice journaling. Generate 7 unique topics (one for each day of the week) that are interesting, thought-provoking, and suitable for 2-3 minute voice recordings.

Topics should:
- Be natural, conversational prompts (not formal interview questions)
- Be engaging and relatable to everyday life
- Encourage personal reflection or storytelling
- Vary in type (memories, opinions, experiences, observations, hobbies, daily life, plans)
- Help build natural conversational fluency
- Feel universal and broadly relatable

AVOID: Formal interview-style questions, structured "Tell me about a time when..." prompts

Respond with a JSON object containing a 'topics' array with 7 topic strings."""
    
    VOICE_JOURNALING_TOPICS_USER_PROMPT = """Generate 7 unique daily topics for this week's voice journaling practice. Each topic should be interesting, natural, and help practice speaking spontaneously for 2-3 minutes about everyday life. Make them varied and engaging - not formal interview questions. Respond in JSON format: {"topics": ["topic 1", "topic 2", ..., "topic 7"]}"""
    
    WEEKLY_PROMPT_WORDS_SYSTEM_PROMPT_EXTRA = """You are generating 5 words for weekly speaking prompt practice focused on product management. These words will be used in a 3-5 minute speaking practice session where the learner discusses PM concepts in professional contexts.

IMPORTANT: Avoid words that have appeared in 3 or more consecutive recent weeks. It's fine if a word from 5 weeks ago comes back (that shows it's important), but avoid words that appeared week after week recently. The system will provide you with words that appeared 3+ consecutive weeks to avoid.

The words should be relevant to product management topics like: strategy, user experience, metrics, prioritization, problem-solving, stakeholder management, product development, user research, business analysis, or product launches.

For each word, provide: the word, its part of speech, and a brief context hint (one sentence) that helps the learner understand how to use it in PM contexts.
Respond with a JSON object containing a 'words' array, where each word is an object with 'word', 'part_of_speech', and 'hint' fields."""
    
    WEEKLY_PROMPT_WORDS_USER_PROMPT = """Generate 5 product management-related words for this week's speaking prompt practice. These words should help build PM vocabulary and can be naturally incorporated into a 3-5 minute speaking practice about PM concepts. Focus on practical PM terminology that will be useful in professional contexts. Respond in JSON format: {"words": [{"word": "...", "part_of_speech": "...", "hint": "..."}, ...]}"""
    
    SHADOWING_SCRIPT_SYSTEM_PROMPT_EXTRA = """You are creating a 7-10 minute shadowing practice script in the style of a TED talk. This script will be converted to audio and used for daily shadowing practice to improve pronunciation, rhythm, and natural speaking flow.

CRITICAL LENGTH REQUIREMENT: The script MUST be approximately 875-1,250 words (spoken at ~125 words per minute = 7-10 minutes). This is essential - generate the full length script, not a summary.

IMPORTANT: Choose a DIFFERENT topic from previous weeks to avoid repetition and provide variety. The system will provide you with summaries of previous scripts' topics.

The script should be in TED talk style:
- Engaging, inspiring, and thought-provoking
- Educational content that teaches something interesting
- Clear structure with an introduction, main points, and conclusion
- Natural, conversational delivery style
- Varied sentence structures for engaging rhythm
- Topics can be: science, psychology, productivity, innovation, personal growth, technology, culture, history, creativity, or any inspiring educational content

The script should be interesting to speak aloud, have natural pauses, and help build speaking confidence through engaging content."""
    
    SHADOWING_SCRIPT_USER_PROMPT = """Generate a complete 7-10 minute shadowing practice script in TED talk style.

CRITICAL REQUIREMENTS:
1. The script MUST be 875-1,250 words (count your words to ensure this)
2. This is a FULL script, not a summary or outline
3. At 125 words per minute, 875-1,250 words = 7-10 minutes of speaking time
4. Write the complete script word-for-word as it would be spoken

SCRIPT STRUCTURE (write each section fully):
- Introduction (100-150 words): Hook the audience, introduce the topic
- Main Point 1 (200-250 words): First key idea with examples/stories
- Main Point 2 (200-250 words): Second key idea with examples/stories  
- Main Point 3 (200-250 words): Third key idea with examples/stories
- Main Point 4 (150-200 words): Fourth key idea with examples/stories (optional but encouraged for 10-minute talks)
- Conclusion (100-150 words): Summarize and leave with a thought

TED TALK STYLE:
- Inspiring, educational, engaging content
- Personal stories or examples to illustrate points
- Clear narrative flow
- Varied sentence lengths for natural rhythm
- Topic: Choose from science, psychology, productivity, innovation, personal growth, technology, culture, history, creativity

IMPORTANT: Count your words as you write. The final script MUST be between 875-1,250 words. Do not stop until you reach this word count. Write the full script now."""
    
    WEEKLY_PROMPT_SYSTEM_PROMPT_EXTRA = """You are creating a weekly speaking prompt for product management practice (3-5 minutes of speaking). This is part of a 24-month preparation journey where the learner is building PM thinking and communication skills.

IMPORTANT: Create a DIFFERENT prompt from previous weeks to avoid repetition and provide variety. The system will provide you with previous prompts to avoid.

The prompt should:
- Be related to product management (strategy, user experience, metrics, prioritization, problem-solving, stakeholder management, product development, etc.)
- Encourage 3-5 minutes of thoughtful speaking
- Help develop PM thinking and communication skills
- Be a prompt that helps practice articulating PM concepts in a natural, conversational way
- Be engaging, thought-provoking, and progressively challenging
- Help build confidence in speaking about PM topics
- Cover new PM topics or approaches that haven't been covered in recent weeks

The learner is in Phase 1 (0-6 months), focusing on building consistency and natural delivery, so the prompt should be approachable yet substantive."""
    
    WEEKLY_PROMPT_USER_PROMPT = """Generate a weekly speaking prompt for product management practice. The prompt should: 1) Be related to product management (e.g., strategy, user experience, metrics, prioritization, problem-solving, stakeholder management, product launches, user research), 2) Encourage 3-5 minutes of speaking, 3) Help develop PM thinking and communication skills for professional contexts, 4) Be something that helps practice speaking about PM concepts in a natural way, 5) Be engaging and thought-provoking, appropriate for Phase 1 learning (building foundation and fluency). Make it something the learner can speak about naturally while incorporating PM vocabulary and concepts."""
    
    # Default prompts for best answer generation (shadowing mode)
    WEEKLY_PROMPT_BEST_ANSWER_SYSTEM_PROMPT = """You are an expert product management career coach who has fully internalized the content from "Cracking the PM Interview" and Lewis Lin's interview answer techniques."""
    WEEKLY_PROMPT_BEST_ANSWER_USER_PROMPT = """Generate a complete "best answer" script for a product management interview question. This answer will be used for shadowing practice (3-5 minutes of speaking). Length: 375-625 words, pure speech text only."""
    WEEKLY_PROMPT_BEST_ANSWER_HINTS_PROMPT = """Generate explanatory hints for the best answer script. These hints should explain each part/section of the generated best answer."""

try:
    from openai import OpenAI
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False
    print("Warning: openai package not installed. Run: pip install openai")


def get_openai_client():
    """Get OpenAI client instance."""
    if not OPENAI_AVAILABLE:
        raise ImportError("openai package is not installed. Install it with: pip install openai")
    
    if not OPENAI_API_KEY:
        raise ValueError("OPENAI_API_KEY is not set in environment variables or .env file")
    
    return OpenAI(api_key=OPENAI_API_KEY)


def load_pm_questions(csv_path='references/pm_questions.csv', sample_size=20):
    """
    Load PM questions from CSV file (for internal reference only).
    
    Args:
        csv_path: Path to the CSV file containing PM questions
        sample_size: Number of questions to return (for context in prompts)
    
    Returns:
        List of question strings, or empty list if file not found
    """
    questions = []
    try:
        if not os.path.exists(csv_path):
            print(f"Warning: {csv_path} not found. Weekly prompts will not reference PM questions.")
            return questions
        
        with open(csv_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for row in reader:
                question = row.get('What was the interview question?', '').strip()
                if question:
                    questions.append(question)
        
        # Return a random sample to keep prompts varied
        if len(questions) > sample_size:
            questions = random.sample(questions, min(sample_size, len(questions)))
        
        return questions
    except Exception as e:
        print(f"Error loading PM questions from {csv_path}: {e}")
        return questions


def generate_voice_journaling_topics(previous_topics=None, regenerate=False):
    """
    Generate 7 daily topics for voice journaling (one for each day of the week).
    
    Args:
        previous_topics: List of previous topics from past weeks to avoid repetition
        regenerate: If True, indicates this is a regeneration and should generate different content
    
    Returns:
        List of 7 daily topics (strings).
    """
    try:
        client = get_openai_client()
        
        user_prompt = VOICE_JOURNALING_TOPICS_USER_PROMPT
        
        if previous_topics:
            user_prompt += f"\n\nIMPORTANT: Avoid these topics from recent weeks: {', '.join(previous_topics[:20])}. Generate fresh, different topics."
        
        if regenerate:
            user_prompt += "\n\nNOTE: This is a regeneration - please generate COMPLETELY DIFFERENT topics from what was previously generated for this week."
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": VOICE_JOURNALING_TOPICS_SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.9 if regenerate else 0.7,
            response_format={"type": "json_object"}
        )
        
        result = response.choices[0].message.content
        import json
        data = json.loads(result)
        
        if 'topics' in data and isinstance(data['topics'], list) and len(data['topics']) == 7:
            return data['topics']
        else:
            # Fallback topics
            return [
                "Describe your ideal weekend",
                "Talk about a recent challenge you overcame",
                "What motivates you to keep learning?",
                "Share a memorable experience from the past year",
                "Discuss a book, movie, or article that impacted you",
                "What are you grateful for today?",
                "Describe your goals for the next month"
            ]
    except Exception as e:
        print(f"Error generating voice journaling topics: {e}")
        return [
            "Describe your ideal weekend",
            "Talk about a recent challenge you overcame",
            "What motivates you to keep learning?",
            "Share a memorable experience from the past year",
            "Discuss a book, movie, or article that impacted you",
            "What are you grateful for today?",
            "Describe your goals for the next month"
        ]

def generate_weekly_prompt_words(previous_words=None, regenerate=False):
    """
    Generate 5 words to include in weekly speaking prompt practice.
    
    Args:
        previous_words: List of previous words from past weeks to avoid repetition
        regenerate: If True, indicates this is a regeneration and should generate different content
    
    Returns:
        List of 5 words with context/hints, focused on product management vocabulary.
    """
    try:
        client = get_openai_client()
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": f"""{BACKGROUND_CONTEXT}

{WEEKLY_PROMPT_WORDS_SYSTEM_PROMPT_EXTRA}"""
                },
                {
                    "role": "user",
                    "content": WEEKLY_PROMPT_WORDS_USER_PROMPT + (
                        f"\n\nIMPORTANT: Avoid these words that appeared in 3+ consecutive recent weeks: {previous_words}. "
                        f"It's fine if a word from 5+ weeks ago comes back (shows importance), but avoid words that appeared week after week recently."
                        if previous_words else ""
                    ) + (
                        "\n\nNOTE: This is a regeneration - please generate COMPLETELY DIFFERENT words from what was previously generated for this week."
                        if regenerate else ""
                    )
                }
            ],
            temperature=0.9 if regenerate else 0.7,  # Higher temperature for regeneration
            response_format={"type": "json_object"}
        )
        
        result = response.choices[0].message.content
        import json
        data = json.loads(result)
        
        # Extract words from the response
        if 'words' in data and isinstance(data['words'], list):
            return data['words']
        elif isinstance(data, list):
            return data
        else:
            # Fallback: try to find any array in the response
            for key in data:
                if isinstance(data[key], list):
                    return data[key]
            return []
            
    except Exception as e:
        print(f"Error generating weekly prompt words: {e}")
        # Return fallback words (5 PM-related words)
        return [
            {"word": "prioritize", "part_of_speech": "verb", "hint": "Decide what to focus on first"},
            {"word": "stakeholder", "part_of_speech": "noun", "hint": "Someone with interest in the product"},
            {"word": "metric", "part_of_speech": "noun", "hint": "A way to measure success"},
            {"word": "iterate", "part_of_speech": "verb", "hint": "Improve through repeated cycles"},
            {"word": "align", "part_of_speech": "verb", "hint": "Get everyone on the same page"}
        ]


def generate_shadowing_scripts(previous_scripts=None, regenerate=False):
    """
    Generate TWO shadowing practice scripts (500-700 words each).
    
    Args:
        previous_scripts: List of previous script summaries/topics from past weeks to avoid repetition
        regenerate: If True, indicates this is a regeneration and should generate different content
    
    Returns:
        Dictionary with 'script1' and 'script2' keys, each containing 500-700 words.
    """
    try:
        client = get_openai_client()
        
        # Build regeneration-specific instructions
        regeneration_note = ""
        if regenerate and previous_scripts:
            # Get the current script that needs to be replaced
            current_script_summary = previous_scripts[0] if previous_scripts else ""
            regeneration_note = f"""

CRITICAL: This is a REGENERATION request. The current scripts for this week start with:
"{current_script_summary[:200]}..."

You MUST generate COMPLETELY DIFFERENT scripts with:
1. DIFFERENT topics for both scripts (not daily routines, productivity, or communication)
2. DIFFERENT openings (not "Welcome to today's shadowing practice")
3. DIFFERENT structure and content
4. Fresh, unique TED talk-style topics (science, psychology, innovation, culture, history, creativity, etc.)
5. Each script should be 500-700 words with COMPLETELY NEW content

DO NOT use any part of the existing scripts. Create something entirely new."""
        elif regenerate:
            regeneration_note = "\n\nCRITICAL: This is a REGENERATION request. Generate COMPLETELY NEW and DIFFERENT scripts. Do not reuse any previous content."
        
        # Build user prompt based on prompts.py, adapted for TWO scripts
        previous_topics_note = ""
        if previous_scripts:
            previous_topics_note = f"\n\nIMPORTANT: Choose DIFFERENT topics from previous weeks. Previous topics: {', '.join([s[:50] + '...' if len(s) > 50 else s for s in (previous_scripts or [])[:3]])}"
        
        # Use the prompt from prompts.py but adapted for TWO scripts
        user_prompt = f"""Generate TWO separate 4-5.5 minute TED talk-style speech scripts for shadowing practice with a presentation / public speech feeling.

CRITICAL REQUIREMENTS FOR BOTH SCRIPTS:
- Generate TWO DIFFERENT scripts (Script 1 and Script 2)
- Each script MUST be 500-700 words (count your words carefully)
- Each script should cover a DIFFERENT topic
- At 125 words per minute, 500-700 words = 4-5.5 minutes of speaking time
- Format: PURE SPEECH TEXT ONLY - NO labels, NO markers like "Title:", "Introduction (125 words):", "[Introduction]", "[Main Point 1]", "[Conclusion]", etc.
- Write ONLY what would be spoken out loud, word-for-word
- The script should start directly with the first spoken words and flow naturally to the last words
- Someone reading this should be able to speak it out loud exactly as written without encountering any formatting elements
- Each script can be a complete segment OR a snippet/excerpt from a longer talk - it doesn't need to follow a strict structure with opening, main points, and conclusion
- The script should feel like a natural, coherent segment of speech that flows smoothly

TED TALK STYLE:
- Inspiring, educational, engaging content
- Personal stories or examples to illustrate points
- Clear narrative flow
- Varied sentence lengths for natural rhythm
- Topics: Choose from science, psychology, innovation, culture, history, creativity, human behavior, society, environment, technology (general topics - NOT career/job specific)

{previous_topics_note}
{regeneration_note}

Example of what NOT to do:
Title: The Power of Mindfulness
Introduction (125 words):
Imagine waking up one morning...
Main Point 1:
Let's start with...

Example of what TO do:
Imagine waking up one morning to find... (and continue with the full speech, flowing naturally from start to finish)

Return your response in JSON format:
{{
    "script1": "Full text of first script (500-700 words, pure speech text only, no labels or markers)...",
    "script2": "Full text of second script (500-700 words, pure speech text only, no labels or markers)..."
}}"""
            
        response = client.chat.completions.create(
            model="gpt-4o",  # Using gpt-4o for better length compliance
            messages=[
                {
                    "role": "system",
                    "content": f"""{BACKGROUND_CONTEXT}

{SHADOWING_SCRIPT_SYSTEM_PROMPT_EXTRA}

You are creating TWO shadowing practice scripts in TED talk style. Each script should be 500-700 words (4-5.5 minutes of speaking time). NEVER include section markers or labels - only the speech text itself."""
                },
                {
                    "role": "user",
                    "content": user_prompt
                }
            ],
            temperature=0.95 if regenerate else 0.8,
            max_tokens=4000,  # Enough for two 500-700 word scripts
            response_format={"type": "json_object"}
        )
        
        result = response.choices[0].message.content.strip()
        import json
        data = json.loads(result)
        
        # Validate the response
        if 'script1' not in data or 'script2' not in data:
            raise Exception("Response missing script1 or script2")
        
        script1 = data['script1'].strip()
        script2 = data['script2'].strip()
        
        # Check word counts
        word_count1 = len(script1.split())
        word_count2 = len(script2.split())
        
        if word_count1 < 400 or word_count2 < 400:
            print(f"Warning: Scripts may be too short (Script 1: {word_count1} words, Script 2: {word_count2} words). Target: 500-700 words each.")
        elif word_count1 > 800 or word_count2 > 800:
            print(f"Warning: Scripts may be too long (Script 1: {word_count1} words, Script 2: {word_count2} words). Target: 500-700 words each.")
        else:
            print(f"Generated scripts - Script 1: {word_count1} words, Script 2: {word_count2} words")
        
        return {
            'script1': script1,
            'script2': script2
        }
        
    except Exception as e:
        print(f"Error generating shadowing scripts: {e}")
        import traceback
        traceback.print_exc()
        # Re-raise the exception instead of returning fallback during regeneration
        if regenerate:
            raise Exception(f"Failed to regenerate shadowing scripts: {str(e)}")
        # Return fallback scripts only for initial generation
        return {
            'script1': """Have you ever wondered why some mornings feel energizing while others feel like a struggle? The secret lies in understanding your body's natural rhythms.

Our bodies operate on a 24-hour cycle called the circadian rhythm. This internal clock influences everything from our energy levels to our ability to focus. When we work with our natural rhythms instead of against them, we unlock incredible potential.

Morning people aren't just born that way - they've learned to sync their activities with their peak performance times. Research shows that our cognitive abilities peak at different times throughout the day. For most people, complex problem-solving is easier in the morning, while creative thinking flourishes in the afternoon.

The key is to identify your own patterns. Track your energy levels for a week. Notice when you feel most alert and when you naturally slow down. Then, schedule your most important tasks during your peak hours.

Small changes make a big difference. Going to bed and waking up at consistent times helps regulate your circadian rhythm. Even on weekends, try to maintain your schedule within an hour of your weekday routine.

Understanding your body's natural rhythms isn't about forcing yourself to be productive 24/7. It's about working smarter, not harder. When you align your life with your biology, productivity becomes effortless.""",
            'script2': """Think about the last time you learned something completely new. Maybe it was a language, a musical instrument, or a new skill at work. Did you feel frustrated at first? That's actually a good sign.

Neuroscientists have discovered something fascinating about how our brains learn. When we struggle with new information, our brains create stronger neural pathways. This process, called "productive failure," is essential for deep learning.

The comfortable feeling of mastery is wonderful, but it doesn't trigger the same growth. When we're challenged and make mistakes, our brains release chemicals that strengthen memory formation. Each error is actually building a more robust understanding.

This is why effective practice looks different from what many people expect. It's not about mindless repetition. It's about working at the edge of your abilities, where success isn't guaranteed. Musicians call this the "learning zone" - not so easy that it's boring, not so hard that it's impossible.

The next time you feel frustrated while learning something new, remember this: your brain is doing exactly what it needs to do. That discomfort means you're growing. Embrace the struggle, learn from your mistakes, and trust the process.

Real learning happens when we're willing to be uncomfortable. The path to mastery is paved with productive failures."""
        }


# Keep the old function for backwards compatibility
def generate_shadowing_script(previous_scripts=None, regenerate=False):
    """
    Generate a single shadowing practice script (backwards compatibility).
    Now returns script1 from the two-script generation.
    
    Args:
        previous_scripts: List of previous script summaries/topics from past weeks to avoid repetition
        regenerate: If True, indicates this is a regeneration and should generate different content
    
    Returns:
        Script text (500-700 words).
    """
    result = generate_shadowing_scripts(previous_scripts, regenerate)
    return result['script1']


def generate_weekly_prompt(previous_prompts=None, regenerate=False, week_key=None):
    """
    Generate a weekly speaking prompt for practice.
    
    Args:
        previous_prompts: List of previous prompts from past weeks to avoid repetition
        regenerate: If True, indicates this is a regeneration and should generate different content
        week_key: Week key (e.g., '2025-W45') - optional, for future use
    
    Returns:
        A prompt that encourages 3-5 minutes of speaking.
        Focused on product management topics for professional communication practice.
        Returns a string prompt.
    """
    try:
        client = get_openai_client()
        
        # Check if PM questions database exists (for context reference)
        pm_questions_exist = os.path.exists('references/pm_questions.csv')
        
        # Build context about PM questions (reference, not copy)
        pm_questions_context = ""
        if pm_questions_exist:
            pm_questions_context = """

IMPORTANT CONTEXT - Reference PM Question Database:
The learner has a database of PM questions covering various topics like: product strategy, metrics and success measurement, prioritization, user research, stakeholder management, product design, execution, A/B testing, and behavioral scenarios.

When generating the weekly prompt and hints, reference these question types conceptually:
- Align the prompt topics with the PM domains covered in those questions
- Structure hints to help build thinking patterns and frameworks that would be useful for similar professional scenarios
- Focus on PM concepts and skills that appear frequently in professional PM contexts
- Help develop structured thinking and communication for professional settings

DO NOT copy or directly use questions from the database. Instead, create original prompts that reference and align with those question types and domains."""
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": f"""{BACKGROUND_CONTEXT}

{WEEKLY_PROMPT_SYSTEM_PROMPT_EXTRA}"""
                },
                {
                    "role": "user",
                    "content": WEEKLY_PROMPT_USER_PROMPT + pm_questions_context + (
                        f"\n\nIMPORTANT: Create a DIFFERENT prompt from previous weeks. "
                        f"Previous prompts ({len(previous_prompts or [])}): {', '.join([p[:80] + '...' if len(p) > 80 else p for p in (previous_prompts or [])[:3]])}. "
                        f"Choose a new PM topic or approach that hasn't been covered recently."
                        if previous_prompts else ""
                    ) + (
                        "\n\nNOTE: This is a regeneration - please generate a COMPLETELY DIFFERENT prompt from what was previously generated for this week."
                        if regenerate else ""
                    )
                }
            ],
            temperature=0.9 if regenerate else 0.8,  # Higher temperature for regeneration
            max_tokens=500  # Increased to allow for hints
        )
        
        prompt = response.choices[0].message.content.strip()
        return prompt
        
    except Exception as e:
        print(f"Error generating weekly prompt: {e}")
        # Return fallback prompts related to product management
        fallback_prompts = [
            "Imagine you're a product manager launching a new feature. Walk through how you would gather user feedback and decide whether to iterate or pivot.",
            "Explain how you would balance user needs with business goals when making product decisions. Use a specific example to illustrate your thinking.",
            "Describe your approach to prioritizing features when resources are limited. What framework would you use and why?",
            "Talk about a product you use daily and analyze what makes it successful from a product management perspective.",
            "Imagine you're explaining a complex product decision to stakeholders with different priorities. How would you structure your explanation?"
        ]
        import random
        return random.choice(fallback_prompts)


def get_resume_context():
    """
    Read and extract text from the resume file to provide context for generating relatable stories.
    
    Returns:
        String containing resume text, or empty string if resume not found or can't be read
    """
    resume_path = 'references/Resume_Hebinna Jeong_062325 (Dolby).docx'
    
    if not os.path.exists(resume_path):
        # Try to find any resume file
        resume_files = [f for f in os.listdir('references') if 'resume' in f.lower() and (f.endswith('.docx') or f.endswith('.pdf'))]
        if resume_files:
            resume_path = os.path.join('references', resume_files[0])
        else:
            return ""
    
    if not DOCX_AVAILABLE:
        return ""
    
    try:
        doc = Document(resume_path)
        resume_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                resume_text.append(paragraph.text.strip())
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    resume_text.append(" | ".join(row_text))
        
        return "\n".join(resume_text)
    except Exception as e:
        print(f"Error reading resume: {e}")
        return ""


def generate_weekly_prompt_best_answer(previous_prompts=None, regenerate=False, week_key=None):
    """
    Generate a weekly prompt with best answer script for shadowing mode.
    
    Args:
        previous_prompts: List of previous prompts from past weeks to avoid repetition
        regenerate: If True, indicates this is a regeneration and should generate different content
        week_key: Week key (e.g., '2025-W45')
    
    Returns:
        Dictionary with:
        - 'prompt': The PM interview question/prompt
        - 'best_answer_script': The complete best answer script (375-625 words)
        - 'best_answer_hints': Explanatory hints for each part of the answer
    """
    try:
        client = get_openai_client()
        
        # Check if PM questions database exists (for context reference)
        pm_questions_exist = os.path.exists('references/pm_questions.csv')
        
        # Build context about PM questions (reference, not copy)
        pm_questions_context = ""
        if pm_questions_exist:
            pm_questions_context = """

IMPORTANT CONTEXT - Reference PM Question Database:
The learner has a database of PM questions covering various topics like: product strategy, metrics and success measurement, prioritization, user research, stakeholder management, product design, execution, A/B testing, and behavioral scenarios.

When generating the prompt and best answer, reference these question types conceptually:
- Align the prompt topics with the PM domains covered in those questions
- Structure the best answer to demonstrate strong PM thinking and frameworks
- Focus on PM concepts and skills that appear frequently in PM interviews
- Show structured thinking and communication patterns useful for interviews"""
        
        # First, generate the prompt/question
        prompt_response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": f"""{BACKGROUND_CONTEXT}

{WEEKLY_PROMPT_SYSTEM_PROMPT_EXTRA}"""
                },
                {
                    "role": "user",
                    "content": """Generate a product management interview question/prompt. This will be used to generate a best answer for shadowing practice.

The prompt should:
- Focus on a specific PM topic or skill (user research, prioritization, metrics, stakeholder management, product strategy, etc.)
- Be a clear, specific interview question that can be answered in 3-5 minutes
- Be suitable for generating a strong, structured best answer
- Reference PM interview question types (product strategy, metrics, prioritization, user research, stakeholder management, product design, execution, A/B testing, behavioral scenarios)

Return ONLY the question/prompt itself, no hints or additional text. Start directly with the question.""" + pm_questions_context + (
                        f"\n\nIMPORTANT: Create a DIFFERENT prompt from previous weeks. "
                        f"Previous prompts ({len(previous_prompts or [])}): {', '.join([p[:80] + '...' if len(p) > 80 else p for p in (previous_prompts or [])[:3]])}. "
                        f"Choose a new PM topic or approach that hasn't been covered recently."
                        if previous_prompts else ""
                    ) + (
                        "\n\nNOTE: This is a regeneration - please generate a COMPLETELY DIFFERENT prompt from what was previously generated for this week."
                        if regenerate else ""
                    )
                }
            ],
            temperature=0.9 if regenerate else 0.8,
            max_tokens=200
        )
        
        prompt = prompt_response.choices[0].message.content.strip()
        
        # Get resume context for relatable stories
        resume_context = get_resume_context()
        resume_context_section = ""
        if resume_context:
            # Limit resume context to avoid token limits (keep it concise)
            resume_summary = resume_context[:2000]  # First 2000 chars should be enough
            resume_context_section = f"""

IMPORTANT - LEARNER'S BACKGROUND CONTEXT:
The learner's resume/work history is provided below. When creating stories and examples, make them relatable to this background. The stories don't have to be exact matches, but should be in domains/industries/contexts the learner would understand and relate to based on their experience.

Resume/Background:
{resume_summary}

Use this context to create stories that feel authentic and relatable to the learner's background. The examples should align with domains, technologies, or work contexts similar to what's in their resume."""
        
        # Now generate the best answer script
        best_answer_response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": f"""{BACKGROUND_CONTEXT}

{WEEKLY_PROMPT_BEST_ANSWER_SYSTEM_PROMPT}"""
                },
                {
                    "role": "user",
                    "content": f"""{WEEKLY_PROMPT_BEST_ANSWER_USER_PROMPT}

The PM interview question is:
{prompt}
{resume_context_section}

Generate the best answer script now. Remember: 375-625 words, pure speech text only, no labels or markers. Include specific PM details (features, campaigns, MVP composition, prioritization decisions) and make the story relatable to the learner's background."""
                }
            ],
            temperature=0.7,
            max_tokens=1000
        )
        
        best_answer_script = best_answer_response.choices[0].message.content.strip()
        
        # Finally, generate explanatory hints
        hints_response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": f"""{BACKGROUND_CONTEXT}

{WEEKLY_PROMPT_BEST_ANSWER_SYSTEM_PROMPT}"""
                },
                {
                    "role": "user",
                    "content": f"""{WEEKLY_PROMPT_BEST_ANSWER_HINTS_PROMPT}

The PM interview question is:
{prompt}

The best answer script is:
{best_answer_script[:500]}...

Generate explanatory hints for this best answer now."""
                }
            ],
            temperature=0.7,
            max_tokens=500
        )
        
        best_answer_hints = hints_response.choices[0].message.content.strip()
        
        return {
            'prompt': prompt,
            'best_answer_script': best_answer_script,
            'best_answer_hints': best_answer_hints
        }
        
    except Exception as e:
        print(f"Error generating weekly prompt best answer: {e}")
        import traceback
        traceback.print_exc()
        # Return fallback
        fallback_prompt = "How would you prioritize features when resources are limited?"
        fallback_answer = """When prioritizing features with limited resources, I would use a structured framework to ensure we're making data-driven decisions. First, I'd gather input from key stakeholders including product, engineering, design, and business teams to understand all potential features and their requirements. Then I'd evaluate each feature against multiple criteria: user impact, business value, technical complexity, and strategic alignment. I'd use a framework like RICE scoring which considers reach, impact, confidence, and effort. Features with high user impact and business value that align with our strategic goals would rank highest. I'd also consider dependencies between features and potential risks. After scoring, I'd present the prioritized list to stakeholders for discussion and alignment, ensuring everyone understands the rationale. This approach ensures we're investing resources in features that deliver the most value while maintaining transparency and buy-in from the team."""
        fallback_hints = """1. Framework Introduction: This opening establishes a structured approach, showing systematic thinking which is valued in PM interviews.\n\n2. Stakeholder Input: This section demonstrates understanding of cross-functional collaboration, a key PM skill.\n\n3. Evaluation Criteria: Shows knowledge of prioritization frameworks (RICE) and multi-dimensional thinking.\n\n4. Strategic Alignment: Demonstrates ability to connect features to business goals.\n\n5. Communication and Buy-in: Shows understanding of stakeholder management and transparency."""
        return {
            'prompt': fallback_prompt,
            'best_answer_script': fallback_answer,
            'best_answer_hints': fallback_hints
        }


def generate_shadowing_audio_openai(script_text, output_path=None, speed=1.0, voice="onyx", return_timestamps=False):
    """
    Generate audio from script text using OpenAI TTS.
    
    Args:
        script_text: The script text to convert to audio
        output_path: Optional path to save audio file (if None, returns audio bytes)
        speed: Audio speed (default 1.0, range 0.25-4.0)
        voice: OpenAI voice (default "onyx", options: alloy, echo, fable, onyx, nova, shimmer)
        return_timestamps: If True, returns tuple of (audio, paragraph_timestamps)
    
    Returns:
        If return_timestamps is True: tuple of (audio_path/bytes, paragraph_timestamps)
        Otherwise: audio_path or audio_bytes
        On error, returns None.
        
        paragraph_timestamps format: [{"paragraph_index": 0, "start_time": 0.0, "text": "..."}]
    """
    try:
        client = get_openai_client()
        
        # Split by paragraphs to track timestamps
        import re
        paragraphs = re.split(r'\n\n+', script_text)
        paragraphs = [p.strip() for p in paragraphs if p.strip()]
        
        # Generate timestamps based on estimated durations
        paragraph_timestamps = []
        current_time = 0.0
        
        for i, paragraph in enumerate(paragraphs):
            paragraph_timestamps.append({
                "paragraph_index": i,
                "start_time": current_time,
                "text": paragraph[:100] + "..." if len(paragraph) > 100 else paragraph
            })
            
            # Estimate duration based on text length and speed
            # Average speaking rate: ~125 words per minute
            word_count = len(paragraph.split())
            duration_seconds = (word_count / 125.0) * 60.0 / speed
            current_time += duration_seconds
        
        # Use OpenAI TTS API
        # Model options: tts-1 (faster, less quality) or tts-1-hd (higher quality)
        # Voice options: alloy, echo, fable, onyx, nova, shimmer
        # Validate voice parameter
        valid_voices = ['alloy', 'echo', 'fable', 'onyx', 'nova', 'shimmer']
        if voice not in valid_voices:
            print(f"Invalid voice '{voice}', defaulting to 'onyx'")
            voice = 'onyx'
        
        response = client.audio.speech.create(
            model="tts-1",  # Using standard quality for speed
            voice=voice,
            input=script_text,
            speed=speed
        )
        
        # Get audio bytes
        audio_bytes = response.content
        
        if output_path:
            # Save to file
            import os
            dir_path = os.path.dirname(output_path)
            if dir_path:
                os.makedirs(dir_path, exist_ok=True)
            
            with open(output_path, 'wb') as f:
                f.write(audio_bytes)
            
            if return_timestamps:
                return (output_path, paragraph_timestamps)
            return output_path
        else:
            # Return bytes
            if return_timestamps:
                return (audio_bytes, paragraph_timestamps)
            return audio_bytes
            
    except Exception as e:
        print(f"Error generating audio with OpenAI: {e}")
        import traceback
        traceback.print_exc()
        return None


def generate_shadowing_audio_openai_for_week(script_text, week_key, speed=1.0, voice="onyx", return_timestamps=False):
    """
    Generate OpenAI audio file for a specific week and save it.
    
    Args:
        script_text: The script text to convert to audio
        week_key: Week key (e.g., '2024-W45')
        speed: Audio speed (default 1.0, range 0.25-4.0)
        voice: OpenAI voice (default "onyx", options: alloy, echo, fable, onyx, nova, shimmer)
        return_timestamps: If True, returns tuple of (audio_url, paragraph_timestamps)
    
    Returns:
        If return_timestamps is True: tuple of (audio_url, paragraph_timestamps)
        Otherwise: Relative path to the audio file (e.g., 'audio/week_2024-W45_openai.mp3') or None on error
    """
    if not script_text or not script_text.strip():
        print("No script text provided")
        return None
    
    # Create audio directory if it doesn't exist
    import os
    audio_dir = os.path.join('static', 'audio')
    os.makedirs(audio_dir, exist_ok=True)
    
    # Generate filename based on week (with openai suffix to distinguish from TypeCast version)
    # OpenAI TTS returns mp3 format
    audio_filename = f"week_{week_key}_openai.mp3"
    output_path = os.path.join(audio_dir, audio_filename)
    
    # Generate audio with specified speed, voice, and get timestamps
    result = generate_shadowing_audio_openai(script_text, output_path=output_path, speed=speed, voice=voice, return_timestamps=return_timestamps)
    
    if result:
        if return_timestamps:
            # result is (output_path, paragraph_timestamps)
            audio_path, timestamps = result
            # Return relative path from static folder for serving
            return (f"audio/{audio_filename}", timestamps)
        else:
            # Return relative path from static folder for serving
            return f"audio/{audio_filename}"
    else:
        return None

